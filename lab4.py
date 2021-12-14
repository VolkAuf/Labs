import docx
import re
from bitarray import bitarray
import time


# Формулка для расчета количества количества хэшфункций: k=(m/n)ln 2 m-размер массива, n-размер данных
# Сайт с хэш-функциями http://vak.ru/doku.php/proj/hash/sources

class BloomFilter(object):

    def __init__(self, size):
        self.size = size
        self.number_of_functions = 10
        self.bloom_filter = bitarray(self.size)
        self.bloom_filter.setall(0)

    def _hash_djb2(self, s):
        hash = 5381
        for x in s:
            hash = ((hash << 5) + hash) + ord(x)
        return hash % self.size

    def _hash_djb2_2(self, s):
        hash = 2139062143
        for x in s:
            hash = ((hash << 7) + hash) + ord(x)
        return hash % self.size

    def _hash_char_code(self, s):
        sum = 0
        for pos in range(len(s)):
            sum = sum + ord(s[pos]) * pos
        return sum % self.size

    def _rs_hash(self, s):
        b = 378551
        a = 63689
        hash = 0
        i = 0
        for x in s:
            hash = hash * a + ord(x)
            a *= b
        return hash % self.size

    def _ly_hash(self, s):
        hash = 0
        for x in s:
            hash = (hash * 1664525) + ord(x) + 1013904223
        return hash % self.size

    def _rot13_hash(self, s):
        hash = 0
        for x in s:
            hash += ord(x)
            hash -= (hash << 13) | (hash >> 19)
        return hash % self.size

    def _faq6_hash(self, s):
        hash = 0
        for x in s:
            hash += ord(x)
            hash += (hash << 13)
            hash ^= (hash >> 6)
        hash += (hash << 3)
        hash ^= (hash >> 1)
        hash += (hash << 15)
        return hash % self.size

    def _hash_h37(self, s):
        hash = 2139062143

        for x in s:
            hash = 37 * hash + ord(x)
        return hash % self.size

    def _js_hash(self, s):
        hash = 1315423911
        for x in s:
            hash ^= ((hash << 5) + ord(x) + (hash >> 2))
        return hash % self.size

    def _dec_hash(self, s):
        hash = len(s)
        for x in s:
            hash = ((hash << 5) ^ (hash >> 27)) ^ ord(x)
        return hash % self.size

    def add_to_filter(self, string):
        self.bloom_filter[self._hash_djb2(string)] = 1
        self.bloom_filter[self._hash_djb2_2(string)] = 1
        self.bloom_filter[self._hash_char_code(string)] = 1
        self.bloom_filter[self._rs_hash(string)] = 1
        self.bloom_filter[self._ly_hash(string)] = 1
        self.bloom_filter[self._rot13_hash(string)] = 1
        self.bloom_filter[self._faq6_hash(string)] = 1
        self.bloom_filter[self._hash_h37(string)] = 1
        self.bloom_filter[self._js_hash(string)] = 1
        self.bloom_filter[self._dec_hash(string)] = 1

    def check_in_filter(self, string):
        if self.bloom_filter[self._hash_djb2(string)] == 0:
            return False
        elif self.bloom_filter[self._hash_djb2_2(string)] == 0:
            return False
        elif self.bloom_filter[self._hash_char_code(string)] == 0:
            return False
        elif self.bloom_filter[self._rs_hash(string)] == 0:
            return False
        elif self.bloom_filter[self._ly_hash(string)] == 0:
            return False
        elif self.bloom_filter[self._rot13_hash(string)] == 0:
            return False
        elif self.bloom_filter[self._faq6_hash(string)] == 0:
            return False
        elif self.bloom_filter[self._hash_h37(string)] == 0:
            return False
        elif self.bloom_filter[self._js_hash(string)] == 0:
            return False
        elif self.bloom_filter[self._dec_hash(string)] == 0:
            return False
        return True

    def merge_filter(self, filter):
        i = 0
        while i < self.size:
            self.bloom_filter[i] = self.bloom_filter[i] or filter.bloom_filter[i]
            i += 1

    def get_fullness(self):
        count = 0
        for item in self.bloom_filter:
            if item == 1:
                count += 1
        return count / self.size


def read_doc(doc, filter, set):
    for text in doc.paragraphs:
        without_sep = re.sub(r'[^\w\s]', '', str(text.text))
        without_sep_arr = without_sep.split()
        set.update(without_sep_arr)
        for word in without_sep_arr:
            filter.add_to_filter(word)


def check_word(word, filter1, filter2, tolstoy):
    print('Есть ли слово в фильтре 1: ', filter1.check_in_filter(word.lower()))
    print('Есть ли слово в фильтре 2: ', filter2.check_in_filter(word.lower()))
    print('Добавлялось ли слово в фильтр 1: ', word.lower() in tolstoy)


def start():
    print("Введите размер фильтра")
    val = int(input())


    filter1 = BloomFilter(val)
    filter2 = BloomFilter(val)

    tolstoy = set()
    tolkin = set()

    read_doc(docx.Document("voina-i-mir.docx"), filter1, tolstoy)
    read_doc(docx.Document("tolkin.docx"), filter2, tolkin)

    if filter1.get_fullness() >= 0.8:
        raise ValueError("Фильтр переполнен")

    size = len(tolkin)
    noise = 0

    for word in tolkin:
        if (not (word in tolstoy)) and (filter1.check_in_filter(word)):
            noise += 1

    print("Слов в фильтре: ", len(tolstoy))
    print("Частота ложноположительных срабатываний: ", noise / size * 100, "%")

    print("Поиск слова Саруман в фильтре 1:")
    start_time = time.perf_counter_ns()
    print(filter1.check_in_filter("Хоббиты"))
    first_time =(time.perf_counter_ns() - start_time)
    print("Перебор нашел слово за - " + str(first_time/1000000) + " мс")

    filter1.merge_filter(filter2)
    print("Поиск слова Саруман в фильтре 1 после объединения с фильтром 2:")
    print(filter1.check_in_filter("Хоббиты"))
    print("Перебор нашел слово за - " + str((time.perf_counter_ns() - first_time)/1000000000) + " с")
